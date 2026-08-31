from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


IMAGE_EXTENSIONS = {
    "png",
    "jpg",
    "jpeg",
}

DOCUMENT_EXTENSIONS = {
    "pdf",
    "docx",
    "txt",
    "md",
    "py",
    "json",
    "csv",
}

SUPPORTED_FILE_EXTENSIONS = (
    IMAGE_EXTENSIONS
    | DOCUMENT_EXTENSIONS
)

# 防止超大文档直接把模型上下文撑爆。
MAX_DOCUMENT_CHARS = 60_000


def get_file_extension(
    uploaded_file: Any,
) -> str:
    """返回小写文件扩展名，不包含点号。"""

    name = getattr(
        uploaded_file,
        "name",
        "",
    )

    return (
        Path(name)
        .suffix
        .lower()
        .lstrip(".")
    )


def is_image_file(
    uploaded_file: Any,
) -> bool:
    return (
        get_file_extension(uploaded_file)
        in IMAGE_EXTENSIONS
    )


def is_document_file(
    uploaded_file: Any,
) -> bool:
    return (
        get_file_extension(uploaded_file)
        in DOCUMENT_EXTENSIONS
    )


def _decode_text(data: bytes) -> str:
    """兼容常见中英文文本编码。"""

    encodings = (
        "utf-8-sig",
        "utf-8",
        "gb18030",
        "big5",
    )

    for encoding in encodings:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode(
        "utf-8",
        errors="replace",
    )


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(
        BytesIO(data)
    )

    pages: list[str] = []

    for index, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            text = (
                page.extract_text()
                or ""
            )
        except Exception:
            text = ""

        text = text.strip()

        if text:
            pages.append(
                f"--- Page {index} ---\n"
                f"{text}"
            )

    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    document = Document(
        BytesIO(data)
    )

    parts: list[str] = []

    # 正文段落
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # 表格
    for table_index, table in enumerate(
        document.tables,
        start=1,
    ):
        rows: list[str] = []

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            rows.append(
                " | ".join(cells)
            )

        if rows:
            parts.append(
                f"\n--- Table "
                f"{table_index} ---\n"
                + "\n".join(rows)
            )

    return "\n\n".join(parts)


def _extract_json(data: bytes) -> str:
    raw_text = _decode_text(data)

    try:
        parsed = json.loads(
            raw_text
        )

        return json.dumps(
            parsed,
            ensure_ascii=False,
            indent=2,
        )

    except Exception:
        return raw_text


def extract_document_text(
    uploaded_file: Any,
) -> dict[str, Any]:
    """
    从 Megor 支持的文档中提取文字。

    返回：
    {
        "filename": ...,
        "extension": ...,
        "text": ...,
        "truncated": bool,
    }
    """

    filename = getattr(
        uploaded_file,
        "name",
        "uploaded_file",
    )

    extension = get_file_extension(
        uploaded_file
    )

    if extension not in DOCUMENT_EXTENSIONS:
        raise ValueError(
            f"Unsupported document type: "
            f"{extension}"
        )

    data = uploaded_file.getvalue()

    if not data:
        raise ValueError(
            "The uploaded file is empty."
        )

    if extension == "pdf":
        text = _extract_pdf(data)

    elif extension == "docx":
        text = _extract_docx(data)

    elif extension == "json":
        text = _extract_json(data)

    else:
        # txt / md / py / csv
        text = _decode_text(data)

    text = text.strip()

    if not text:
        if extension == "pdf":
            raise ValueError(
                "No readable text was found "
                "in this PDF. It may be a "
                "scanned/image-only PDF."
            )

        raise ValueError(
            "No readable text was found "
            "in this file."
        )

    truncated = False

    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[
            :MAX_DOCUMENT_CHARS
        ]

        truncated = True

    return {
        "filename": filename,
        "extension": extension,
        "text": text,
        "truncated": truncated,
    }


def build_document_prompt(
    uploaded_file: Any,
    user_prompt: str,
) -> str:
    """
    把文件内容统一转成普通文本上下文，
    这样所有 Megor 模型都可以使用。
    """

    result = extract_document_text(
        uploaded_file
    )

    filename = result["filename"]
    text = result["text"]
    truncated = result["truncated"]

    question = (
        user_prompt.strip()
        if user_prompt
        else (
            "Please analyze and summarize "
            "this document."
        )
    )

    truncation_notice = ""

    if truncated:
        truncation_notice = (
            "\n\n[Notice: The document was "
            "too long, so only the first "
            f"{MAX_DOCUMENT_CHARS:,} "
            "characters are included.]"
        )

    return (
        f"The user uploaded a file named "
        f'"{filename}".\n\n'
        f"===== FILE CONTENT =====\n"
        f"{text}"
        f"{truncation_notice}\n"
        f"===== END FILE CONTENT =====\n\n"
        f"User request:\n"
        f"{question}"
    )